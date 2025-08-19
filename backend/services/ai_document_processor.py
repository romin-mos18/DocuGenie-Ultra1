#!/usr/bin/env python3
"""
AI Document Processor - Complete AI-powered document analysis pipeline
Extracts content using AI tools, analyzes it, and provides intelligent labeling
"""
import os
import logging
import json
from typing import Dict, List, Optional, Any
from datetime import datetime
from pathlib import Path

# AI Processing imports
try:
    from PIL import Image
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    print("⚠️ Tesseract OCR not available. Install: pip install pytesseract pillow")

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    print("⚠️ PyMuPDF not available. Install: pip install PyMuPDF")

try:
    from docx import Document as DocxDocument
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    print("⚠️ python-docx not available. Install: pip install python-docx")

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    print("⚠️ pandas not available. Install: pip install pandas")

# Text analysis imports
import re
from collections import Counter

logger = logging.getLogger(__name__)

class AIDocumentProcessor:
    """Complete AI-powered document processing system"""
    
    def __init__(self):
        """Initialize AI document processor"""
        self.processing_stats = {
            "total_processed": 0,
            "successful": 0,
            "failed": 0,
            "last_processed": None
        }
        
        # Document type patterns with enhanced AI analysis
        self.document_patterns = {
            "medical_report": {
                "keywords": [
                    "patient", "diagnosis", "treatment", "symptoms", "medical", "report",
                    "physician", "doctor", "clinic", "hospital", "examination", "findings",
                    "consultation", "assessment", "clinical", "history", "condition", "therapeutic",
                    "prognosis", "evaluation", "progress note", "chief complaint", "vital signs",
                    "physical examination", "assessment and plan", "impression", "recommendation"
                ],
                "patterns": [
                    r"medical\s+report",
                    r"patient\s*:\s*[A-Z][a-z]+",
                    r"diagnosis\s*:\s*",
                    r"chief\s+complaint",
                    r"physical\s+examination",
                    r"assessment\s+and\s+plan",
                    r"vital\s+signs",
                    r"dr\.\s+[A-Z][a-z]+",
                    r"physician\s*:\s*"
                ],
                "confidence_boost": 0.3
            },
            
            "lab_result": {
                "keywords": [
                    "laboratory", "test", "result", "blood", "urine", "analysis", "lab",
                    "biochemistry", "hematology", "microbiology", "pathology", "values",
                    "specimen", "glucose", "hemoglobin", "cholesterol", "culture", "biopsy",
                    "normal range", "abnormal", "reference range", "lab report", "test results"
                ],
                "patterns": [
                    r"laboratory\s+report",
                    r"lab\s+report",
                    r"test\s+results?",
                    r"\d+\s*mg/dL",
                    r"\d+\s*g/dL",
                    r"normal\s*:\s*\d+",
                    r"reference\s+range",
                    r"specimen\s*:\s*",
                    r"glucose\s*:\s*\d+",
                    r"hemoglobin\s*:\s*\d+"
                ],
                "confidence_boost": 0.25
            },
            
            "prescription": {
                "keywords": [
                    "prescription", "medication", "drug", "dosage", "pharmacy", "rx",
                    "tablet", "capsule", "injection", "refill", "prescribed", "take",
                    "dose", "daily", "twice", "instructions", "quantity", "mg", "mcg"
                ],
                "patterns": [
                    r"prescription",
                    r"medication\s*:\s*",
                    r"take\s+\d+",
                    r"\d+\s*mg",
                    r"\d+\s*mcg",
                    r"tablet[s]?",
                    r"capsule[s]?",
                    r"refill[s]?\s*:\s*\d+",
                    r"quantity\s*:\s*\d+",
                    r"pharmacy\s+instructions"
                ],
                "confidence_boost": 0.28
            },
            
            "clinical_trial": {
                "keywords": [
                    "clinical", "trial", "study", "protocol", "investigation", "research",
                    "participant", "informed consent", "phase", "randomized", "placebo",
                    "efficacy", "safety", "endpoint", "inclusion", "exclusion", "criteria"
                ],
                "patterns": [
                    r"clinical\s+trial",
                    r"research\s+protocol",
                    r"phase\s+[I1-4IV]+",
                    r"informed\s+consent",
                    r"inclusion\s+criteria",
                    r"exclusion\s+criteria",
                    r"primary\s+endpoint",
                    r"randomized\s+controlled"
                ],
                "confidence_boost": 0.22
            },
            
            "insurance": {
                "keywords": [
                    "insurance", "policy", "coverage", "claim", "benefits", "premium",
                    "deductible", "copay", "provider", "network", "authorization",
                    "member", "subscriber", "plan", "coverage"
                ],
                "patterns": [
                    r"insurance\s+policy",
                    r"policy\s+number",
                    r"member\s+id",
                    r"subscriber\s*:\s*",
                    r"coverage\s+details",
                    r"premium\s*:\s*\$?\d+",
                    r"deductible\s*:\s*\$?\d+",
                    r"copay\s*:\s*\$?\d+"
                ],
                "confidence_boost": 0.2
            },
            
            "billing": {
                "keywords": [
                    "bill", "invoice", "payment", "charge", "cost", "fee", "amount",
                    "balance", "statement", "account", "financial", "total", "due"
                ],
                "patterns": [
                    r"invoice\s+number",
                    r"bill\s+to\s*:",
                    r"total\s+amount\s*:\s*\$?\d+",
                    r"balance\s+due\s*:\s*\$?\d+",
                    r"payment\s+due",
                    r"\$\s*\d+\.\d{2}",
                    r"account\s+number"
                ],
                "confidence_boost": 0.18
            }
        }
        
        logger.info("✅ AI Document Processor initialized")
    
    def process_document(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """
        Complete AI-powered document processing pipeline
        
        Args:
            file_path: Path to the document file
            file_type: Type of file (pdf, docx, jpg, etc.)
            
        Returns:
            Complete processing results with AI analysis and labeling
        """
        start_time = datetime.now()
        
        try:
            logger.info(f"🤖 Starting AI document processing: {file_path}")
            
            # Step 1: AI Content Extraction
            extraction_result = self._extract_content_with_ai(file_path, file_type)
            
            if not extraction_result["success"]:
                return self._create_error_result("Content extraction failed", start_time, extraction_result["error"])
            
            extracted_text = extraction_result["text"]
            
            # Step 2: AI Content Analysis
            analysis_result = self._analyze_content_with_ai(extracted_text)
            
            # Step 3: AI-powered Document Labeling
            labeling_result = self._label_document_with_ai(extracted_text, analysis_result)
            
            # Step 4: Generate AI Summary
            summary_result = self._generate_ai_summary(extracted_text, labeling_result["document_type"])
            
            # Step 5: Extract Key Information using AI
            key_info_result = self._extract_key_information(extracted_text, labeling_result["document_type"])
            
            # Combine all results
            processing_time = (datetime.now() - start_time).total_seconds()
            
            final_result = {
                "success": True,
                "processing_method": "ai_powered",
                "processing_time": processing_time,
                "timestamp": datetime.now().isoformat(),
                
                # Extraction Results
                "extraction": {
                    "success": True,
                    "text": extracted_text,
                    "word_count": len(extracted_text.split()),
                    "character_count": len(extracted_text),
                    "extraction_method": extraction_result["method"],
                    "confidence": extraction_result["confidence"]
                },
                
                # AI Analysis Results
                "analysis": analysis_result,
                
                # AI Labeling Results
                "classification": {
                    "document_type": labeling_result["document_type"],
                    "confidence": labeling_result["confidence"],
                    "ai_reasoning": labeling_result["reasoning"],
                    "alternative_types": labeling_result.get("alternatives", []),
                    "success": True
                },
                
                # AI Summary
                "summary": summary_result,
                
                # Key Information
                "key_information": key_info_result,
                
                # Processing Stats
                "processing_stats": {
                    "extraction_time": extraction_result.get("processing_time", 0),
                    "analysis_time": analysis_result.get("processing_time", 0),
                    "labeling_time": labeling_result.get("processing_time", 0),
                    "total_time": processing_time
                }
            }
            
            # Update stats
            self.processing_stats["total_processed"] += 1
            self.processing_stats["successful"] += 1
            self.processing_stats["last_processed"] = datetime.now().isoformat()
            
            logger.info(f"✅ AI processing completed: {labeling_result['document_type']} ({labeling_result['confidence']:.1%})")
            
            return final_result
            
        except Exception as e:
            logger.error(f"❌ AI document processing failed: {e}")
            self.processing_stats["total_processed"] += 1
            self.processing_stats["failed"] += 1
            return self._create_error_result(str(e), start_time)
    
    def _extract_content_with_ai(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Extract content using AI-powered tools"""
        start_time = datetime.now()
        
        try:
            extracted_text = ""
            method = "unknown"
            confidence = 0.0
            
            # PDF Processing with AI
            if file_type.lower() == 'pdf':
                if PYMUPDF_AVAILABLE:
                    extracted_text, confidence = self._extract_pdf_with_ai(file_path)
                    method = "pymupdf_ai"
                else:
                    raise Exception("PDF processing not available")
            
            # Image Processing with OCR AI
            elif file_type.lower() in ['jpg', 'jpeg', 'png', 'bmp', 'tiff', 'gif']:
                if TESSERACT_AVAILABLE:
                    extracted_text, confidence = self._extract_image_with_ocr_ai(file_path)
                    method = "tesseract_ocr_ai"
                else:
                    raise Exception("OCR processing not available")
            
            # Word Document Processing
            elif file_type.lower() in ['docx', 'doc']:
                if PYTHON_DOCX_AVAILABLE:
                    extracted_text, confidence = self._extract_docx_with_ai(file_path)
                    method = "docx_ai"
                else:
                    raise Exception("DOCX processing not available")
            
            # Excel/CSV Processing
            elif file_type.lower() in ['xlsx', 'xls', 'csv']:
                if PANDAS_AVAILABLE:
                    extracted_text, confidence = self._extract_excel_with_ai(file_path)
                    method = "excel_ai"
                else:
                    raise Exception("Excel processing not available")
            
            # Plain Text Processing
            elif file_type.lower() in ['txt', 'md', 'rtf']:
                extracted_text, confidence = self._extract_text_with_ai(file_path)
                method = "text_ai"
            
            else:
                raise Exception(f"Unsupported file type: {file_type}")
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return {
                "success": True,
                "text": extracted_text,
                "method": method,
                "confidence": confidence,
                "processing_time": processing_time,
                "word_count": len(extracted_text.split()),
                "character_count": len(extracted_text)
            }
            
        except Exception as e:
            logger.error(f"❌ Content extraction failed: {e}")
            return {
                "success": False,
                "error": str(e),
                "text": "",
                "method": "failed",
                "confidence": 0.0
            }
    
    def _extract_pdf_with_ai(self, pdf_path: str) -> tuple:
        """Extract text from PDF using AI-enhanced methods"""
        try:
            doc = fitz.open(pdf_path)
            all_text = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Extract text
                text = page.get_text()
                if text.strip():
                    all_text.append(text)
                
                # For pages with little text, try OCR on images
                if len(text.strip()) < 50:
                    try:
                        pix = page.get_pixmap()
                        img_data = pix.tobytes("ppm")
                        # Could add OCR processing here
                    except:
                        pass
            
            doc.close()
            
            full_text = "\n".join(all_text)
            confidence = 0.95 if len(full_text) > 100 else 0.7
            
            return full_text, confidence
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return "", 0.0
    
    def _extract_image_with_ocr_ai(self, image_path: str) -> tuple:
        """Extract text from images using OCR AI"""
        try:
            # Open and preprocess image
            image = Image.open(image_path)
            
            # Enhance image for better OCR
            image = image.convert('RGB')
            
            # Extract text using Tesseract
            extracted_text = pytesseract.image_to_string(image, config='--psm 6')
            
            # Calculate confidence based on text quality
            confidence = self._calculate_ocr_confidence(extracted_text)
            
            return extracted_text, confidence
            
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            return "", 0.0
    
    def _extract_docx_with_ai(self, docx_path: str) -> tuple:
        """Extract text from DOCX with AI-enhanced processing"""
        try:
            doc = DocxDocument(docx_path)
            all_text = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    all_text.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        all_text.append(" | ".join(row_text))
            
            full_text = "\n".join(all_text)
            confidence = 0.98  # DOCX extraction is very reliable
            
            return full_text, confidence
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return "", 0.0
    
    def _extract_excel_with_ai(self, excel_path: str) -> tuple:
        """Extract data from Excel with AI processing"""
        try:
            if excel_path.endswith('.csv'):
                df = pd.read_csv(excel_path)
                text_content = df.to_string(index=False)
            else:
                # Read all sheets
                excel_file = pd.ExcelFile(excel_path)
                all_data = []
                
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(excel_path, sheet_name=sheet_name)
                    sheet_data = f"Sheet: {sheet_name}\n{df.to_string(index=False)}"
                    all_data.append(sheet_data)
                
                text_content = "\n\n".join(all_data)
            
            confidence = 0.95
            return text_content, confidence
            
        except Exception as e:
            logger.error(f"Excel extraction failed: {e}")
            return "", 0.0
    
    def _extract_text_with_ai(self, text_path: str) -> tuple:
        """Extract and enhance plain text"""
        try:
            with open(text_path, 'r', encoding='utf-8', errors='ignore') as file:
                text_content = file.read()
            
            confidence = 1.0  # Plain text is 100% reliable
            return text_content, confidence
            
        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            return "", 0.0
    
    def _analyze_content_with_ai(self, text: str) -> Dict[str, Any]:
        """Analyze document content using AI techniques"""
        start_time = datetime.now()
        
        try:
            analysis = {
                "success": True,
                "processing_time": 0,
                "content_analysis": {},
                "structural_analysis": {},
                "semantic_analysis": {}
            }
            
            # Content Analysis
            analysis["content_analysis"] = {
                "word_count": len(text.split()),
                "character_count": len(text),
                "sentence_count": len(re.split(r'[.!?]+', text)),
                "paragraph_count": len([p for p in text.split('\n\n') if p.strip()]),
                "avg_sentence_length": len(text.split()) / max(len(re.split(r'[.!?]+', text)), 1),
                "readability_score": self._calculate_readability(text)
            }
            
            # Structural Analysis
            analysis["structural_analysis"] = {
                "has_headers": bool(re.search(r'^[A-Z][A-Z\s]+:?$', text, re.MULTILINE)),
                "has_lists": bool(re.search(r'^\s*[\d\-\*]\s+', text, re.MULTILINE)),
                "has_tables": bool(re.search(r'\|.*\|', text)),
                "has_dates": bool(re.search(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b', text)),
                "has_numbers": bool(re.search(r'\b\d+\b', text)),
                "has_medical_terms": self._has_medical_terminology(text),
                "document_structure": self._analyze_document_structure(text)
            }
            
            # Semantic Analysis
            analysis["semantic_analysis"] = {
                "key_terms": self._extract_key_terms(text),
                "named_entities": self._extract_named_entities(text),
                "topic_indicators": self._identify_topic_indicators(text),
                "language_complexity": self._analyze_language_complexity(text)
            }
            
            analysis["processing_time"] = (datetime.now() - start_time).total_seconds()
            
            return analysis
            
        except Exception as e:
            logger.error(f"Content analysis failed: {e}")
            return {
                "success": False,
                "error": str(e),
                "processing_time": (datetime.now() - start_time).total_seconds()
            }
    
    def _label_document_with_ai(self, text: str, analysis: Dict) -> Dict[str, Any]:
        """Use AI analysis to intelligently label the document"""
        start_time = datetime.now()
        
        try:
            scores = {}
            reasoning = {}
            
            # Analyze against each document type
            for doc_type, patterns in self.document_patterns.items():
                score = 0
                reasons = []
                
                # Keyword matching with AI enhancement
                keyword_matches = 0
                for keyword in patterns["keywords"]:
                    if keyword.lower() in text.lower():
                        keyword_matches += 1
                        reasons.append(f"Found keyword: '{keyword}'")
                
                keyword_score = keyword_matches / len(patterns["keywords"])
                score += keyword_score * 0.4
                
                # Pattern matching with regex
                pattern_matches = 0
                for pattern in patterns["patterns"]:
                    if re.search(pattern, text, re.IGNORECASE):
                        pattern_matches += 1
                        reasons.append(f"Matched pattern: '{pattern}'")
                
                pattern_score = pattern_matches / len(patterns["patterns"])
                score += pattern_score * 0.4
                
                # Structural analysis boost
                if analysis["success"] and "structural_analysis" in analysis:
                    struct = analysis["structural_analysis"]
                    
                    if doc_type == "medical_report":
                        if struct.get("has_medical_terms", False):
                            score += 0.1
                            reasons.append("Contains medical terminology")
                        if struct.get("has_dates", False):
                            score += 0.05
                            reasons.append("Contains dates")
                    
                    elif doc_type == "lab_result":
                        if struct.get("has_numbers", False):
                            score += 0.1
                            reasons.append("Contains numerical values")
                        if struct.get("has_tables", False):
                            score += 0.05
                            reasons.append("Contains tabular data")
                    
                    elif doc_type == "prescription":
                        if re.search(r'\d+\s*mg\b', text, re.IGNORECASE):
                            score += 0.1
                            reasons.append("Contains dosage information")
                
                # Apply confidence boost
                if score > 0:
                    score += patterns["confidence_boost"] * score
                
                scores[doc_type] = min(score, 1.0)
                reasoning[doc_type] = reasons
            
            # Determine best match
            if scores:
                best_type = max(scores, key=scores.get)
                best_score = scores[best_type]
                
                # Get alternatives
                sorted_scores = sorted(scores.items(), key=lambda x: x[1], reverse=True)
                alternatives = [
                    {"type": doc_type, "confidence": score, "reasoning": reasoning[doc_type][:3]}
                    for doc_type, score in sorted_scores[1:4]
                    if score > 0.1
                ]
                
                # Final confidence adjustment
                if best_score < 0.3:
                    best_type = "other"
                    best_score = 0.5
                
                processing_time = (datetime.now() - start_time).total_seconds()
                
                return {
                    "success": True,
                    "document_type": best_type,
                    "confidence": best_score,
                    "reasoning": reasoning[best_type][:5],
                    "alternatives": alternatives,
                    "all_scores": scores,
                    "processing_time": processing_time
                }
            else:
                return {
                    "success": True,
                    "document_type": "other",
                    "confidence": 0.5,
                    "reasoning": ["No clear type indicators found"],
                    "alternatives": [],
                    "processing_time": (datetime.now() - start_time).total_seconds()
                }
                
        except Exception as e:
            logger.error(f"Document labeling failed: {e}")
            return {
                "success": False,
                "error": str(e),
                "document_type": "other",
                "confidence": 0.0,
                "processing_time": (datetime.now() - start_time).total_seconds()
            }
    
    def _generate_ai_summary(self, text: str, document_type: str) -> Dict[str, Any]:
        """Generate AI-powered document summary"""
        try:
            # Extract first few sentences as summary
            sentences = re.split(r'[.!?]+', text)
            clean_sentences = [s.strip() for s in sentences if s.strip() and len(s.strip()) > 20]
            
            # Create summary based on document type
            if document_type == "medical_report":
                summary_sentences = [s for s in clean_sentences[:5] if any(term in s.lower() for term in ["patient", "diagnosis", "treatment", "examination"])]
            elif document_type == "lab_result":
                summary_sentences = [s for s in clean_sentences[:3] if any(term in s.lower() for term in ["test", "result", "normal", "abnormal"])]
            else:
                summary_sentences = clean_sentences[:3]
            
            if not summary_sentences:
                summary_sentences = clean_sentences[:2]
            
            summary = ". ".join(summary_sentences[:3]) + "." if summary_sentences else "No summary available."
            
            return {
                "success": True,
                "summary": summary,
                "summary_type": "ai_generated",
                "word_count": len(summary.split()),
                "sentence_count": len(summary_sentences)
            }
            
        except Exception as e:
            logger.error(f"Summary generation failed: {e}")
            return {
                "success": False,
                "error": str(e),
                "summary": "Summary generation failed."
            }
    
    def _extract_key_information(self, text: str, document_type: str) -> Dict[str, Any]:
        """Extract key information based on document type"""
        try:
            key_info = {}
            
            # Common extractions
            key_info["dates"] = re.findall(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b', text)
            key_info["names"] = re.findall(r'\b[A-Z][a-z]+ [A-Z][a-z]+\b', text)
            key_info["numbers"] = re.findall(r'\b\d+(?:\.\d+)?\b', text)
            
            # Document-specific extractions
            if document_type == "medical_report":
                key_info["medical_terms"] = [term for term in ["diagnosis", "treatment", "symptoms", "medication"] if term in text.lower()]
                key_info["vitals"] = re.findall(r'BP\s*:?\s*\d+/\d+|HR\s*:?\s*\d+|Temp\s*:?\s*\d+\.?\d*', text, re.IGNORECASE)
                
            elif document_type == "lab_result":
                key_info["lab_values"] = re.findall(r'\b\w+\s*:\s*\d+(?:\.\d+)?\s*\w+/\w+', text)
                key_info["normal_ranges"] = re.findall(r'Normal\s*:\s*[\d\-<>]+', text, re.IGNORECASE)
                
            elif document_type == "prescription":
                key_info["medications"] = re.findall(r'\b[A-Z][a-z]+\s+\d+\s*mg', text)
                key_info["dosages"] = re.findall(r'take\s+\d+.*?daily|twice.*?day', text, re.IGNORECASE)
            
            return {
                "success": True,
                "key_information": key_info,
                "extraction_count": sum(len(v) if isinstance(v, list) else 1 for v in key_info.values())
            }
            
        except Exception as e:
            logger.error(f"Key information extraction failed: {e}")
            return {
                "success": False,
                "error": str(e),
                "key_information": {}
            }
    
    def _calculate_ocr_confidence(self, text: str) -> float:
        """Calculate OCR confidence based on text quality"""
        if not text or len(text.strip()) < 10:
            return 0.1
        
        # Check for common OCR errors
        error_indicators = len(re.findall(r'[^\w\s\.\,\!\?\;\:\-\(\)\'\"\/]', text))
        word_count = len(text.split())
        
        if word_count == 0:
            return 0.1
        
        error_ratio = error_indicators / word_count
        confidence = max(0.3, 1.0 - (error_ratio * 2))
        
        return min(confidence, 0.95)
    
    def _calculate_readability(self, text: str) -> float:
        """Calculate basic readability score"""
        words = text.split()
        sentences = re.split(r'[.!?]+', text)
        
        if not words or not sentences:
            return 0.0
        
        avg_words_per_sentence = len(words) / len(sentences)
        
        # Simple readability approximation
        if avg_words_per_sentence < 15:
            return 0.8  # Easy
        elif avg_words_per_sentence < 25:
            return 0.6  # Medium
        else:
            return 0.4  # Difficult
    
    def _has_medical_terminology(self, text: str) -> bool:
        """Check if text contains medical terminology"""
        medical_terms = [
            "patient", "diagnosis", "treatment", "symptoms", "medication",
            "physician", "doctor", "clinical", "medical", "hospital",
            "examination", "therapy", "prescription", "laboratory"
        ]
        
        text_lower = text.lower()
        return sum(1 for term in medical_terms if term in text_lower) >= 3
    
    def _analyze_document_structure(self, text: str) -> Dict[str, Any]:
        """Analyze document structure"""
        lines = text.split('\n')
        
        return {
            "total_lines": len(lines),
            "empty_lines": len([line for line in lines if not line.strip()]),
            "header_lines": len([line for line in lines if line.isupper() and len(line.split()) <= 5]),
            "indented_lines": len([line for line in lines if line.startswith('  ') or line.startswith('\t')]),
            "numbered_lines": len([line for line in lines if re.match(r'^\s*\d+\.', line)])
        }
    
    def _extract_key_terms(self, text: str) -> List[str]:
        """Extract key terms from text"""
        # Simple keyword extraction
        words = re.findall(r'\b[A-Za-z]{4,}\b', text.lower())
        word_freq = Counter(words)
        
        # Filter out common words
        common_words = {"that", "this", "with", "from", "they", "been", "have", "were", "said", "each", "which", "their", "time", "will", "about", "would", "there", "could", "other", "more", "very", "what", "know", "just", "first", "into", "over", "think", "also", "your", "work", "life", "only", "can", "still", "should", "after", "being", "now", "made", "before", "here", "through", "when", "where", "much", "some", "these", "many", "most", "other", "such", "long", "make", "thing", "see", "him", "two", "more", "go", "no", "way", "may", "say"}
        
        key_terms = [word for word, freq in word_freq.most_common(10) if word not in common_words and freq > 1]
        return key_terms[:5]
    
    def _extract_named_entities(self, text: str) -> Dict[str, List[str]]:
        """Extract named entities"""
        entities = {
            "persons": re.findall(r'\b[A-Z][a-z]+ [A-Z][a-z]+\b', text),
            "dates": re.findall(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b', text),
            "organizations": re.findall(r'\b[A-Z][a-z]+ (?:Hospital|Clinic|Laboratory|University|Company|Corp)\b', text),
            "locations": re.findall(r'\b[A-Z][a-z]+, [A-Z]{2}\b', text)
        }
        
        # Limit results
        for key in entities:
            entities[key] = entities[key][:5]
        
        return entities
    
    def _identify_topic_indicators(self, text: str) -> List[str]:
        """Identify topic indicators"""
        topics = []
        
        if any(term in text.lower() for term in ["patient", "diagnosis", "treatment"]):
            topics.append("healthcare")
        
        if any(term in text.lower() for term in ["laboratory", "test", "result"]):
            topics.append("medical_testing")
        
        if any(term in text.lower() for term in ["prescription", "medication", "drug"]):
            topics.append("pharmaceutical")
        
        if any(term in text.lower() for term in ["insurance", "policy", "coverage"]):
            topics.append("insurance")
        
        if any(term in text.lower() for term in ["bill", "invoice", "payment"]):
            topics.append("billing")
        
        return topics
    
    def _analyze_language_complexity(self, text: str) -> Dict[str, Any]:
        """Analyze language complexity"""
        words = text.split()
        
        if not words:
            return {"complexity": "unknown", "score": 0}
        
        # Calculate average word length
        avg_word_length = sum(len(word) for word in words) / len(words)
        
        # Count complex words (>6 characters)
        complex_words = len([word for word in words if len(word) > 6])
        complexity_ratio = complex_words / len(words)
        
        if avg_word_length < 4.5 and complexity_ratio < 0.3:
            complexity = "simple"
            score = 0.3
        elif avg_word_length < 6 and complexity_ratio < 0.5:
            complexity = "moderate"
            score = 0.6
        else:
            complexity = "complex"
            score = 0.9
        
        return {
            "complexity": complexity,
            "score": score,
            "avg_word_length": avg_word_length,
            "complex_word_ratio": complexity_ratio
        }
    
    def _create_error_result(self, error: str, start_time: datetime, details: str = None) -> Dict[str, Any]:
        """Create error result"""
        processing_time = (datetime.now() - start_time).total_seconds()
        
        return {
            "success": False,
            "error": error,
            "error_details": details,
            "processing_method": "failed",
            "processing_time": processing_time,
            "timestamp": datetime.now().isoformat(),
            "extraction": {"success": False, "text": "", "confidence": 0.0},
            "classification": {"document_type": "unknown", "confidence": 0.0, "success": False}
        }
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Get processing statistics"""
        return {
            "service_name": "AIDocumentProcessor",
            "processing_stats": self.processing_stats,
            "capabilities": {
                "pdf_processing": PYMUPDF_AVAILABLE,
                "ocr_processing": TESSERACT_AVAILABLE,
                "docx_processing": PYTHON_DOCX_AVAILABLE,
                "excel_processing": PANDAS_AVAILABLE,
                "ai_analysis": True,
                "intelligent_labeling": True
            },
            "supported_types": ["pdf", "docx", "doc", "jpg", "jpeg", "png", "bmp", "tiff", "xlsx", "xls", "csv", "txt", "md", "rtf"],
            "document_types": list(self.document_patterns.keys()) + ["other"]
        }
